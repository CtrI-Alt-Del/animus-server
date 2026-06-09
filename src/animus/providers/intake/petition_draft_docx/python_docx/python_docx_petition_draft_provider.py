import re
import shutil
import tempfile
import unicodedata
from pathlib import Path

import docx
from docx.document import Document as DocxDocument

from animus.core.intake.domain.structures.dtos.analysis_document_dto import (
    AnalysisDocumentDto,
)
from animus.core.intake.domain.structures.petition_draft import PetitionDraft
from animus.core.intake.interfaces import PetitionDraftDocxProvider
from animus.core.shared.domain.structures import Datetime, FilePath
from animus.core.storage.interfaces import FileStorageProvider


class PythonDocxPetitionDraftProvider(PetitionDraftDocxProvider):
    _EXPORT_PATH_PREFIX = 'intake/analyses/{analysis_id}/documents/'

    def __init__(self, file_storage_provider: FileStorageProvider) -> None:
        self._file_storage_provider = file_storage_provider

    def export(
        self,
        analysis_id: str,
        analysis_name: str,
        petition_draft: PetitionDraft,
    ) -> AnalysisDocumentDto:
        file_name = self._build_file_name(analysis_name)
        bucket_file_path = FilePath.create(
            self._EXPORT_PATH_PREFIX.format(analysis_id=analysis_id) + file_name
        )
        uploaded_at = Datetime.create_at_now().value.isoformat()

        with tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False
        ) as temporary_file:
            local_file_path = Path(temporary_file.name)

        upload_source_path = Path(bucket_file_path.value)
        upload_source_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            document = self._build_document(petition_draft)
            document.save(str(local_file_path))
            shutil.copy2(local_file_path, upload_source_path)
            self._file_storage_provider.upload_files([bucket_file_path])
        finally:
            local_file_path.unlink(missing_ok=True)
            upload_source_path.unlink(missing_ok=True)

        return AnalysisDocumentDto(
            analysis_id=analysis_id,
            uploaded_at=uploaded_at,
            file_path=bucket_file_path.value,
            name=file_name,
        )

    def _build_document(self, petition_draft: PetitionDraft) -> DocxDocument:
        document = docx.Document()
        self._add_text_section(
            document=document,
            title='Fatos Estruturados',
            content=petition_draft.structured_facts.value,
        )
        self._add_text_section(
            document=document,
            title='Fundamentos Juridicos',
            content=petition_draft.legal_grounds.value,
        )
        self._add_text_section(
            document=document,
            title='Tese Central',
            content=petition_draft.central_thesis.value,
        )
        self._add_list_section(
            document=document,
            title='Pedidos',
            items=[request.value for request in petition_draft.requests],
        )
        self._add_list_section(
            document=document,
            title='Citações de Precedentes',
            items=[
                precedent_citation.value
                for precedent_citation in petition_draft.precedent_citations
            ],
        )

        return document

    def _add_text_section(
        self, document: DocxDocument, title: str, content: str
    ) -> None:
        document.add_heading(title, level=1)
        for paragraph in content.splitlines():
            document.add_paragraph(paragraph)

    def _add_list_section(
        self, document: DocxDocument, title: str, items: list[str]
    ) -> None:
        document.add_heading(title, level=1)
        for item in items:
            document.add_paragraph(item, style='List Bullet')

    def _build_file_name(self, analysis_name: str) -> str:
        normalized_name = unicodedata.normalize('NFKD', analysis_name)
        ascii_name = normalized_name.encode('ascii', 'ignore').decode('ascii')
        sanitized_name = re.sub(r'[^A-Za-z0-9]+', '-', ascii_name).strip('-').lower()

        if sanitized_name == '':
            sanitized_name = 'analise'

        return f'{sanitized_name}-minuta.docx'
